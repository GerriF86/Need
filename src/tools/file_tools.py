<<<<<<< HEAD
import base64
import textract  # use textract to handle multiple file types conveniently

def extract_text_from_file(file_content: str, filename: str) -> str:
    """
    Extract readable text from an uploaded file.
    file_content: base64-encoded content of the file.
    filename: original file name with extension.
    Returns a UTF-8 text string extracted from the file, or an empty string if extraction fails.
    """
    if not file_content or not filename:
        return ""

    # Decode the base64 content to bytes
    try:
        file_bytes = base64.b64decode(file_content)
    except Exception as e:
        print(f"extract_text_from_file: Base64 decode error - {e}")
        return ""

    # Determine file type by extension
    ext = filename.split(".")[-1].lower() if "." in filename else ""

    try:
        # Use textract for simplicity, it supports pdf, docx, txt, etc.
        text = textract.process(file_bytes, extension=ext)
        # textract.process can accept bytes and extension to figure out the parser
        # However, if the above doesn't work as intended, an alternative is writing file_bytes to a temp file and using textract on it.
    except Exception as e:
        print(f"extract_text_from_file: textract failed for {filename} - {e}")
        # Fallback manual handling for basic types
        try:
            if ext == "txt":
                text = file_bytes  # raw bytes of text file (will decode below)
            elif ext == "pdf":
                # Fallback: try PyMuPDF
                import fitz  # PyMuPDF
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                extracted = []
                for page in doc:
                    extracted.append(page.get_text())
                text = "\n".join(extracted).encode('utf-8')
            elif ext in ("doc", "docx", "rtf", "odt"):
                # Fallback: try to use python-docx for docx, or treat others as text
                if ext in ("docx", "odt"):
                    from io import BytesIO
                    from docx import Document
                    doc = Document(BytesIO(file_bytes))
                    full_text = [para.text for para in doc.paragraphs]
                    text = "\n".join(full_text).encode('utf-8')
                else:
                    # .doc or .rtf – not easily handled without external tools. Return empty.
                    text = b""
            else:
                text = b""
        except Exception as e2:
            print(f"extract_text_from_file: fallback extraction failed - {e2}")
            text = b""

    # Decode text bytes to string
    if isinstance(text, bytes):
        try:
            text_str = text.decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"extract_text_from_file: UTF-8 decode error - {e}")
            text_str = text.decode('latin-1', errors='ignore')  # fallback decode
    else:
        # If textract returned a string already
        text_str = str(text)

    # Clean the text: remove excessive whitespace
    text_str = text_str.strip()
    # If text is extremely long, consider truncating or summarizing here (but main summarization is handled upstream).
    return text_str[:10000]  # Return at most 10000 characters to limit size for processing
=======
from __future__ import annotations
from src.utils.tool_registry import tool

@tool
def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extracts text from an uploaded file (PDF, DOCX, or TXT) with improved handling:
    - Uses PyMuPDF (fitz) for PDFs.
    - Wraps DOCX bytes in a BytesIO for python-docx to parse reliably.
    - Maintains paragraph separation rather than merging everything into one line.
    - Basic text cleanup to remove excessive whitespace.

    Raises ValueError if file type is not supported.
    """
    import os
    import fitz
    import docx
    from io import BytesIO

    ext = os.path.splitext(filename)[1].lower()
    text = ""

    if ext == ".pdf":
        # Use PyMuPDF to extract text from PDF
        try:
            with fitz.open(stream=file_content, filetype="pdf") as pdf_doc:
                pages_text = []
                for page in pdf_doc:
                    pages_text.append(page.get_text("text"))  # 'text' preserves layout better than 'blocks'
                text = "\n\n".join(pages_text)
        except Exception as e:
            raise ValueError(f"Error reading PDF: {e}")

    elif ext == ".docx":
        # Wrap bytes in a BytesIO for python-docx
        try:
            file_obj = BytesIO(file_content)
            document = docx.Document(file_obj)
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {e}")

    elif ext == ".txt":
        # Decode as UTF-8 (ignore errors if not valid UTF-8)
        text = file_content.decode("utf-8", errors="ignore")

    else:
        # Unsupported file type
        raise ValueError(f"Unsupported file type: {ext}")

    # --- Text Cleanup ---
    # 1) Convert all kinds of newlines to standard \n
    # 2) Strip extra trailing spaces on each line
    lines = [line.strip() for line in text.splitlines()]
    # 3) Remove empty lines that might result from repeated \n\n
    lines = [ln for ln in lines if ln]
    # Rejoin with one blank line between paragraphs
    cleaned_text = "\n\n".join(lines)

    return cleaned_text
>>>>>>> 07e4a6fe9ee73e2827c50212620e30ba2b339372
